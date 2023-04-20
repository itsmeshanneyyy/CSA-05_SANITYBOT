import pickle
import sklearn
from sanitybot.models import User_health
from flask_login import login_required, current_user
from sanitybot import db
from sanitybot.models import User
from datetime import datetime
from docx import Document
import docx
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING
from io import BytesIO
from datetime import datetime

with open(r'sanitybot\pickles\ppd_decisiontree.pkl', 'rb') as f:
    decisionTree_model = pickle.load(f)
with open(r'sanitybot\pickles\ppd_naiveBayes.pkl', 'rb') as f:
    naiveBayes_model = pickle.load(f)
with open(r'sanitybot\pickles\ppd_lstm.pkl', 'rb') as f:
    lstm_model = pickle.load(f)


def pdd_prediction(input_list):
    pred_result = lstm_model.predict([input_list])[0]
    print(pred_result)

    user = current_user
    user = User.query.filter_by(id=user.id).first()
    user.epds_score = int(pred_result)
    user.date_submitted = datetime.now()
    ppd = ""
    epds_class = ""
    if (pred_result) == 0:
        ppd = """Great news! Based on the responses you provided, it appears that you are not likely experiencing postpartum depression. However, it's always important to prioritize your mental health and seek support if you ever feel overwhelmed or have any concerns. Remember that taking care of yourself is crucial in being the best caregiver for your little one."""
        epds_class = "Depression not likely"
    elif (pred_result) == 1:
        ppd = """Thank you for completing the assessment. Based on your answers, it's possible that you may be experiencing postpartum depression. It's important to discuss your concerns with a healthcare provider to determine the best course of action for your well-being."""
        epds_class = "Depression possible"
    elif (pred_result) == 2:
        ppd = """I'm sorry to hear that the assessment results show a fairly high possibility of postpartum depression. Remember that seeking help is a sign of strength, and there are resources available to support you. """
        epds_class = "Fairly high possibility of depression"
    elif (pred_result) == 3:
        ppd = """I'm sorry to hear that the assessment suggests probable depression. Please know that you're not alone and there is help available. It's important to prioritize your mental health and seek support from healthcare professionals. I'm here to provide resources and support in any way I can."""
        epds_class = "Probable depression"
    else:
        ppd = 'Nothing'

    data = {
        "id": str(user.id),
        "firstname": f"{user.firstname} {user.middlename} {user.lastname}",
        "email": user.email,
        "address": user.address,
        "date_submitted": str(user.date_submitted),
        "epds_score": epds_class,
        "questions": ["QUESTIONS", "Have you been capable of finding humor and laughing about situations?", "Have you anticipated things with pleasure and excitement?", "Have you needlessly held yourself responsible when things didn't go well?", "Have you experienced anxiety or concern without a valid cause?", "Have you experienced fear or panic without a clear or justifiable reason?", "Have things been overwhelming you?", "Have you been so unhappy that you have experienced trouble sleeping?", "Have you experienced feelings of sadness or misery?", "Have you been so unhappy that you have shed tears?", "Have you had thoughts of self-harm?", "EPDS Score"],
        "message": ppd.replace("Thankyou for answering my initial assessment. ", "")
    }

    table = ["Have you been capable of finding humor and laughing about situations?", "Have you anticipated things with pleasure and excitement?", "Have you needlessly held yourself responsible when things didn't go well?", "Have you experienced anxiety or concern without a valid cause?",
             "Have you experienced fear or panic without a clear or justifiable reason?", "Have things been overwhelming you?", "Have you been so unhappy that you have experienced trouble sleeping?", "Have you experienced feelings of sadness or misery?", "Have you been so unhappy that you have shed tears?", "Have you had thoughts of self-harm?"]

    epds = [
        {
            "ANSWER": "",
            0: "As much as I always could",
            1: "Not quite so much now",
            2: "Definitely not so much now",
            3: "Hardly at all",
        },
        {
            "question": "Have you been capable of finding humor and laughing about situations?",
            0: "As much as I always could",
            1: "Not quite so much now",
            2: "Definitely not so much now",
            3: "Hardly at all",
        },
        {
            "question": "Have you anticipated things with pleasure and excitement?",
            0: "As much as I ever did",
            1: "Rather less than I used to",
            2: "Definitely less than I used to",
            3: "Hardly at all",
        },
        {
            "question": "Have you needlessly held yourself responsible when things didn't go well?",
            3: "Yes, most of the time",
            2: "Yes, some of the time",
            1: "Not very often",
            0: "No, Never",
        },
        {
            "question": "Have you experienced anxiety or concern without a valid cause?",
            0: "No, not at all",
            1: "Hardly ever",
            2: "Yes, sometimes",
            3: "Yes, very often",
        },
        {
            "question": "Have you experienced fear or panic without a clear or justifiable reason?",
            3: "Yes, quite a lot",
            2: "Yes, sometimes",
            1: "No, not much",
            0: "No, not at all",
        },
        {
            "question": "Have things been overwhelming you?",
            3: "Yes, most of the time I haven't been able to cope",
            2: "Yes, sometimes I haven't been coping as well as usual",
            1: "No, most of the time I have coped quite well",
            0: "No, I have been coping as well as ever",
        },
        {
            "question": "Have you been so unhappy that you have experienced trouble sleeping?",
            3: "Yes, most of the time",
            2: "Yes, sometimes",
            1: "Not very often",
            0: "No, not at all",
        },
        {
            "question": "Have you experienced feelings of sadness or misery?",
            3: "Yes, most of the time",
            2: "Yes, quite often",
            1: "Not very often",
            0: "No, not at all",
        },
        {
            "question": "Have you been so unhappy that you have shed tears?",
            3: "Yes, most of the time",
            2: "Yes, quite often",
            1: "Only occasionally",
            0: "No, never",
        },
        {
            "question": "Have you had thoughts of self-harm?",
            3: "Yes, quite often",
            2: "Sometimes",
            1: "Hardly ever",
            0: "Never",
        },
        {
            "question": "Have you had thoughts of self-harm?",
            3: "Yes, quite often",
            2: "Sometimes",
            1: "Hardly ever",
            0: "Never",
        },
    ]

    document = Document()
    document.styles['Normal'].paragraph_format.line_spacing = 1.5
    section = document.sections[0]
    section.page_height = Inches(14)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    table = document.add_table(rows=2, cols=3, style="Medium Shading 1")

    datas = ["id", "firstname", "email",
             "address", "date_submitted", "epds_score"]

    table.cell(0, 0).text = 'User ID: ' + data['id']
    table.cell(0, 1).text = 'Name: ' + data['firstname']
    table.cell(0, 2).text = 'Email: ' + data["email"]
    table.cell(1, 0).text = 'Address: ' + data["address"]
    table.cell(1, 1).text = 'Date Submitted: ' + data["date_submitted"]
    table.cell(1, 2).text = 'EPDS Score Interpretation: ' + data["epds_score"]

    # document.add_paragraph()

    # Add a table to the document with 2 columns
    table = document.add_table(
        rows=len(data['questions']), cols=2, style="Medium Shading 1")

    # Access the cells of the first row
    input_list.insert(0, "ANSWER")
    for index, value in enumerate(data['questions']):
        cell_1 = table.cell(index, 0)
        cell_2 = table.cell(index, 1)

        # Add text to the cells
        cell_1.text = value
        if value == "QUESTIONS":
            cell_1.paragraphs[0].runs[0].bold = True
        cell_2.text = f"({input_list[index]}) {epds[index][input_list[index] if input_list[index] in [0,1,2,3] else 0]}"
        if index == 0:
            cell_2.text = "ANSWER"
            cell_2.paragraphs[0].runs[0].bold = True
        elif index == 11:
            cell_2.text = f"{input_list[index]}"

    # table.paragraph_format.space_after = 1

    p = document.add_paragraph(data["message"])
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.space_before = Inches(0.5)

    table = document.add_table(rows=6, cols=3, style="Medium Shading 1")
    table.cell(0, 0).text = "EPDS Score"
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).text = "Interpretation"
    table.cell(0, 1).paragraphs[0].runs[0].bold = True
    table.cell(0, 2).text = "Action"
    table.cell(0, 2).paragraphs[0].runs[0].bold = True

    table.cell(1, 0).text = "Less than 8"
    table.cell(1, 1).text = "Depression not likely"
    table.cell(1, 2).text = "Continue support"
    table.cell(2, 0).text = "9-11"
    table.cell(2, 1).text = "Depression possible"
    table.cell(
        2, 2).text = "Support, re-screen in 2–4 weeks. Consider referral to primary care provide(PCP)."
    table.cell(3, 0).text = "12-13"
    table.cell(3, 1).text = "Fairly high possibility of depression"
    table.cell(3, 2).text = "Monitor, support and offer education. Refer to PCP."
    table.cell(4, 0).text = "14 and higher (positive screen)"
    table.cell(4, 1).text = "Probable depression"
    table.cell(
        4, 2).text = "Diagnostic assessment and treatment by PCP and/or specialist."
    table.cell(
        5, 0).text = "Positive score (1, 2 or 3) on question 10 (suicidality risk)"
    table.cell(5, 1).text = ""
    table.cell(5, 2).text = "Immediate referral is necessary for assessment of suicidal ideation. This is to ensure proper intervention and to consider factors such as plan, history, symptoms, and potential harm."

    new_filename = f"{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
    document.save('./sanitybot/AssessmentResult.docx')
    db.session.commit()

    response = {}

    ppd_result = [ppd, str(pred_result)]

    return ppd_result
